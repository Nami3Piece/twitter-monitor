"""
contract_gen.py — Generic sales contract generator (PDF + Word, CN/EN).
Usage:
    result = generate_contract(params)
    # result = {"cn_pdf": "/tmp/...", "en_pdf": "/tmp/...", "cn_docx": "/tmp/...", "en_docx": "/tmp/..."}
"""

import os
import base64
import tempfile
from datetime import datetime
from typing import Dict, List, Optional


# ── helpers ───────────────────────────────────────────────────────────────────

def _today() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def _fmt_num(n) -> str:
    try:
        return f"{float(n):,.2f}"
    except Exception:
        return str(n)


def _save_logo(logo_b64: str, tmpdir: str) -> Optional[str]:
    """Decode base64 logo and save to temp file. Returns path or None."""
    if not logo_b64:
        return None
    try:
        data = base64.b64decode(logo_b64)
        ext = "png" if data[:8] == b'\x89PNG\r\n\x1a\n' else "jpg"
        path = os.path.join(tmpdir, f"logo.{ext}")
        with open(path, "wb") as f:
            f.write(data)
        return path
    except Exception:
        return None


def _decode_images(products: List[Dict], tmpdir: str) -> List[Dict]:
    """Decode base64 spec_images to temp files, return updated products list."""
    result = []
    for i, p in enumerate(products):
        imgs = p.get("spec_images", [])
        paths = []
        for j, b64 in enumerate(imgs[:3]):
            try:
                data = base64.b64decode(b64)
                ext = "jpg"
                if data[:8] == b'\x89PNG\r\n\x1a\n':
                    ext = "png"
                fpath = os.path.join(tmpdir, f"spec_{i}_{j}.{ext}")
                with open(fpath, "wb") as f:
                    f.write(data)
                paths.append(fpath)
            except Exception:
                pass
        result.append({**p, "_img_paths": paths})
    return result


# ── contract data builder ─────────────────────────────────────────────────────

def _build_data(params: Dict) -> Dict:
    buyer_name    = params.get("buyer_name", "")
    buyer_address = params.get("buyer_address", "")
    buyer_contact = params.get("buyer_contact", "")
    shipping_per  = float(params.get("shipping_per_unit", 50))
    products      = params.get("products", [])
    logo_b64      = params.get("logo_b64", "")
    if not products:
        raise ValueError("At least one product is required")

    # Normalize products
    norm = []
    for p in products:
        qty        = int(p.get("qty", 1))
        unit_price = float(p.get("unit_price", 0))
        norm.append({
            "name":       p.get("name", ""),
            "sku":        p.get("sku", ""),
            "qty":        qty,
            "unit_price": unit_price,
            "subtotal":   qty * unit_price,
            "spec_text":  p.get("spec_text", ""),
            "spec_images": p.get("spec_images", []),
        })

    qty_total      = sum(p["qty"] for p in norm)
    goods_total    = sum(p["subtotal"] for p in norm)
    shipping_total = qty_total * shipping_per
    grand_total    = goods_total + shipping_total

    needs_spec = any(
        len(p.get("spec_text", "")) > 20 or p.get("spec_images")
        for p in norm
    )

    return {
        "buyer_name":     buyer_name,
        "buyer_address":  buyer_address,
        "buyer_contact":  buyer_contact,
        "products":       norm,
        "qty_total":      qty_total,
        "shipping_per":   shipping_per,
        "goods_total":    goods_total,
        "shipping_total": shipping_total,
        "grand_total":    grand_total,
        "needs_spec":     needs_spec,
        "date":           _today(),
        "seller_name":    "Arkreen Network Ltd.",
        "seller_address": "Suite 1, 2nd Floor, The Sotheby Building, Rodney Bay, Gros-Islet, Saint Lucia",
        "seller_contact": "nami3piece@gmail.com",
        "logo_b64":       logo_b64,
    }


# ── PDF page header callback ───────────────────────────────────────────────────

def _make_header_cb(logo_path: Optional[str]):
    """Return an onPage callback that draws logo at top-left of every page."""
    def _draw_header(canvas, doc):
        if not logo_path:
            return
        from reportlab.lib.units import cm
        try:
            canvas.saveState()
            canvas.drawImage(logo_path, x=2*cm, y=doc.pagesize[1] - 1.6*cm,
                             height=1.2*cm, preserveAspectRatio=True, mask='auto')
            canvas.restoreState()
        except Exception:
            pass
    return _draw_header


# ── PDF generator CN ──────────────────────────────────────────────────────────

def _gen_pdf_cn(d: Dict, path: str, tmpdir: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    _cn_font = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/PingFang.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                _cn_font = "CJK"
            except Exception:
                pass
            break

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=3.2*cm, bottomMargin=2*cm)
    title_style = ParagraphStyle("title", fontName=_cn_font, fontSize=16,
                                 spaceAfter=6, alignment=1, leading=22)
    h2_style    = ParagraphStyle("h2", fontName=_cn_font, fontSize=12,
                                 spaceAfter=4, spaceBefore=12, leading=18)
    h3_style    = ParagraphStyle("h3", fontName=_cn_font, fontSize=11,
                                 spaceAfter=4, spaceBefore=8, leading=16)
    body_style  = ParagraphStyle("body", fontName=_cn_font, fontSize=10,
                                 spaceAfter=4, leading=16)

    products_with_imgs = _decode_images(d["products"], tmpdir)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    header_cb = _make_header_cb(logo_path)

    story = []
    story.append(Paragraph("产品销售合同", title_style))
    story.append(Paragraph(f"合同编号：CTR-{d['date'].replace('-','')}-001", body_style))
    story.append(Paragraph(f"签订日期：{d['date']}", body_style))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("一、甲乙双方信息", h2_style))
    story.append(Paragraph(f"<b>甲方（卖方）：</b>{d['seller_name']}", body_style))
    story.append(Paragraph(f"<b>地址：</b>{d['seller_address']}", body_style))
    story.append(Paragraph(f"<b>联系方式：</b>{d['seller_contact']}", body_style))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f"<b>乙方（买方）：</b>{d['buyer_name']}", body_style))
    story.append(Paragraph(f"<b>地址：</b>{d['buyer_address']}", body_style))
    story.append(Paragraph(f"<b>联系方式：</b>{d['buyer_contact']}", body_style))

    story.append(Paragraph("二、产品信息", h2_style))
    tdata = [["产品名称", "产品编号/SKU", "数量", "单价(USD)", "小计(USD)"]]
    for p in d["products"]:
        tdata.append([p["name"], p["sku"], str(p["qty"]),
                      _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])])
    tdata.append(["运费", f"每件 USD {_fmt_num(d['shipping_per'])}", str(d["qty_total"]),
                  _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])])
    tdata.append(["", "", "", "合计", _fmt_num(d["grand_total"])])

    t = Table(tdata, colWidths=[4.5*cm, 3*cm, 2*cm, 3*cm, 3*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,-1), _cn_font),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
        ("ALIGN",      (2,0), (-1,-1), "CENTER"),
        ("BACKGROUND", (0,-1), (-1,-1), colors.HexColor("#f0fdf4")),
        ("ROWBACKGROUNDS", (0,1), (-1,-2), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(t)

    story.append(Paragraph("三、付款方式", h2_style))
    story.append(Paragraph("乙方应在合同签订后 <b>7 个工作日内</b>完成付款，支持 USDT（Polygon 网络）或银行电汇。", body_style))

    story.append(Paragraph("四、交货条款", h2_style))
    story.append(Paragraph("甲方在收到全额货款后 <b>15 个工作日内</b>安排发货，运输方式为国际快递（DHL/FedEx）。", body_style))

    story.append(Paragraph("五、质量保证", h2_style))
    story.append(Paragraph("产品自交货之日起享有 <b>12 个月</b>质量保修，因产品质量问题导致的损失由甲方承担。", body_style))

    story.append(Paragraph("六、违约责任", h2_style))
    story.append(Paragraph("任何一方违约，应向守约方支付合同总金额 <b>10%</b> 的违约金，并赔偿实际损失。", body_style))

    story.append(Paragraph("七、争议解决", h2_style))
    story.append(Paragraph("本合同适用中华人民共和国法律，争议提交合同签订地仲裁委员会仲裁解决。", body_style))

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("八、签署", h2_style))
    sig_data = [
        ["甲方（卖方）签字/盖章", "乙方（买方）签字/盖章"],
        ["\n\n\n", "\n\n\n"],
        [f"日期：{d['date']}", "日期：___________"],
    ]
    sig_t = Table(sig_data, colWidths=[8*cm, 8*cm])
    sig_t.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), _cn_font),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("BOX",      (0,0), (0,-1), 0.5, colors.HexColor("#334155")),
        ("BOX",      (1,0), (1,-1), 0.5, colors.HexColor("#334155")),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(sig_t)

    # Spec section
    if d["needs_spec"]:
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph("附件：产品规格说明 / Product Specifications", h2_style))
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            story.append(Paragraph(label, h3_style))
            if spec_text:
                story.append(Paragraph(spec_text, body_style))
            for img_path in img_paths:
                try:
                    img = RLImage(img_path, width=15*cm)
                    img_w, img_h = img.imageWidth, img.imageHeight
                    if img_h > 0:
                        ratio = img_w / img_h
                        display_h = min(15*cm / ratio, 200)
                        img = RLImage(img_path, width=15*cm, height=display_h)
                    story.append(img)
                    story.append(Spacer(1, 0.3*cm))
                except Exception:
                    pass

    doc.build(story, onFirstPage=header_cb, onLaterPages=header_cb)


# ── PDF generator EN ──────────────────────────────────────────────────────────

def _gen_pdf_en(d: Dict, path: str, tmpdir: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib import colors

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=3.2*cm, bottomMargin=2*cm)
    title_style = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=16,
                                 spaceAfter=6, alignment=1, leading=22)
    h2_style    = ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12,
                                 spaceAfter=4, spaceBefore=12, leading=18)
    h3_style    = ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=11,
                                 spaceAfter=4, spaceBefore=8, leading=16)
    body_style  = ParagraphStyle("body", fontName="Helvetica", fontSize=10,
                                 spaceAfter=4, leading=16)

    products_with_imgs = _decode_images(d["products"], tmpdir)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    header_cb = _make_header_cb(logo_path)

    story = []
    story.append(Paragraph("SALES CONTRACT", title_style))
    story.append(Paragraph(f"Contract No.: CTR-{d['date'].replace('-','')}-001", body_style))
    story.append(Paragraph(f"Date: {d['date']}", body_style))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("1. Parties", h2_style))
    story.append(Paragraph(f"<b>Seller:</b> {d['seller_name']}", body_style))
    story.append(Paragraph(f"<b>Address:</b> {d['seller_address']}", body_style))
    story.append(Paragraph(f"<b>Contact:</b> {d['seller_contact']}", body_style))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f"<b>Buyer:</b> {d['buyer_name']}", body_style))
    story.append(Paragraph(f"<b>Address:</b> {d['buyer_address']}", body_style))
    story.append(Paragraph(f"<b>Contact:</b> {d['buyer_contact']}", body_style))

    story.append(Paragraph("2. Products", h2_style))
    tdata = [["Product Name", "Product No./SKU", "Qty", "Unit Price (USD)", "Subtotal (USD)"]]
    for p in d["products"]:
        tdata.append([p["name"], p["sku"], str(p["qty"]),
                      _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])])
    tdata.append(["Shipping", f"USD {_fmt_num(d['shipping_per'])} / unit", str(d["qty_total"]),
                  _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])])
    tdata.append(["", "", "", "TOTAL", _fmt_num(d["grand_total"])])

    t = Table(tdata, colWidths=[4.5*cm, 3*cm, 2*cm, 3*cm, 3*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTNAME",   (0,1), (-1,-1), "Helvetica"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
        ("ALIGN",      (2,0), (-1,-1), "CENTER"),
        ("BACKGROUND", (0,-1), (-1,-1), colors.HexColor("#f0fdf4")),
        ("FONTNAME",   (3,-1), (-1,-1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0,1), (-1,-2), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(t)

    story.append(Paragraph("3. Payment", h2_style))
    story.append(Paragraph("Buyer shall complete payment within <b>7 business days</b> of contract signing. Accepted methods: USDT (Polygon network) or bank wire transfer.", body_style))

    story.append(Paragraph("4. Delivery", h2_style))
    story.append(Paragraph("Seller shall arrange shipment within <b>15 business days</b> after receipt of full payment via international courier (DHL/FedEx).", body_style))

    story.append(Paragraph("5. Warranty", h2_style))
    story.append(Paragraph("Products carry a <b>12-month</b> quality warranty from the date of delivery. Defects attributable to the Seller will be remedied at Seller's expense.", body_style))

    story.append(Paragraph("6. Liability", h2_style))
    story.append(Paragraph("The breaching party shall pay the non-breaching party a penalty equal to <b>10%</b> of the total contract value, plus actual damages.", body_style))

    story.append(Paragraph("7. Dispute Resolution", h2_style))
    story.append(Paragraph("This contract is governed by the laws of Saint Lucia. Disputes shall be submitted to arbitration at the agreed arbitration body.", body_style))

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("8. Signatures", h2_style))
    sig_data = [
        ["Seller Signature / Stamp", "Buyer Signature / Stamp"],
        ["\n\n\n", "\n\n\n"],
        [f"Date: {d['date']}", "Date: ___________"],
    ]
    sig_t = Table(sig_data, colWidths=[8*cm, 8*cm])
    sig_t.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("BOX",      (0,0), (0,-1), 0.5, colors.HexColor("#334155")),
        ("BOX",      (1,0), (1,-1), 0.5, colors.HexColor("#334155")),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(sig_t)

    if d["needs_spec"]:
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph("Appendix: Product Specifications", h2_style))
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            story.append(Paragraph(label, h3_style))
            if spec_text:
                story.append(Paragraph(spec_text, body_style))
            for img_path in img_paths:
                try:
                    img = RLImage(img_path, width=15*cm)
                    img_w, img_h = img.imageWidth, img.imageHeight
                    if img_h > 0:
                        ratio = img_w / img_h
                        display_h = min(15*cm / ratio, 200)
                        img = RLImage(img_path, width=15*cm, height=display_h)
                    story.append(img)
                    story.append(Spacer(1, 0.3*cm))
                except Exception:
                    pass

    doc.build(story, onFirstPage=header_cb, onLaterPages=header_cb)


# ── Word generator CN ─────────────────────────────────────────────────────────

def _gen_docx_cn(d: Dict, path: str, tmpdir: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    products_with_imgs = _decode_images(d["products"], tmpdir)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    if logo_path:
        try:
            for section in doc.sections:
                header = section.header
                hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = hp.add_run()
                run.add_picture(logo_path, height=Cm(1.2))
        except Exception:
            pass

    title = doc.add_heading("产品销售合同", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"合同编号：CTR-{d['date'].replace('-','')}-001")
    doc.add_paragraph(f"签订日期：{d['date']}")

    doc.add_heading("一、甲乙双方信息", 1)
    doc.add_paragraph(f"甲方（卖方）：{d['seller_name']}")
    doc.add_paragraph(f"地址：{d['seller_address']}")
    doc.add_paragraph(f"联系方式：{d['seller_contact']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"乙方（买方）：{d['buyer_name']}")
    doc.add_paragraph(f"地址：{d['buyer_address']}")
    doc.add_paragraph(f"联系方式：{d['buyer_contact']}")

    doc.add_heading("二、产品信息", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["产品名称", "产品编号/SKU", "数量", "单价(USD)", "小计(USD)"]):
        hdr[i].text = h
    for p in d["products"]:
        row = table.add_row().cells
        for i, val in enumerate([p["name"], p["sku"], str(p["qty"]),
                                  _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])]):
            row[i].text = val
    ship_row = table.add_row().cells
    for i, val in enumerate(["运费", f"每件 USD {_fmt_num(d['shipping_per'])}", str(d["qty_total"]),
                              _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])]):
        ship_row[i].text = val
    total_row = table.add_row().cells
    total_row[3].text = "合计"
    total_row[4].text = _fmt_num(d["grand_total"])

    doc.add_heading("三、付款方式", 1)
    doc.add_paragraph("乙方应在合同签订后 7 个工作日内完成付款，支持 USDT（Polygon 网络）或银行电汇。")

    doc.add_heading("四、交货条款", 1)
    doc.add_paragraph("甲方在收到全额货款后 15 个工作日内安排发货，运输方式为国际快递（DHL/FedEx）。")

    doc.add_heading("五、质量保证", 1)
    doc.add_paragraph("产品自交货之日起享有 12 个月质量保修，因产品质量问题导致的损失由甲方承担。")

    doc.add_heading("六、违约责任", 1)
    doc.add_paragraph("任何一方违约，应向守约方支付合同总金额 10% 的违约金，并赔偿实际损失。")

    doc.add_heading("七、争议解决", 1)
    doc.add_paragraph("本合同适用中华人民共和国法律，争议提交合同签订地仲裁委员会仲裁解决。")

    doc.add_heading("八、签署", 1)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    sig_table.cell(0, 0).text = "甲方（卖方）签字/盖章"
    sig_table.cell(0, 1).text = "乙方（买方）签字/盖章"
    sig_table.cell(1, 0).text = "\n\n"
    sig_table.cell(1, 1).text = "\n\n"
    sig_table.cell(2, 0).text = f"日期：{d['date']}"
    sig_table.cell(2, 1).text = "日期：___________"

    if d["needs_spec"]:
        doc.add_heading("附件：产品规格说明 / Product Specifications", 1)
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            doc.add_heading(label, 2)
            if spec_text:
                doc.add_paragraph(spec_text)
            for img_path in img_paths:
                try:
                    doc.add_picture(img_path, width=Cm(15))
                except Exception:
                    pass

    doc.save(path)


# ── Word generator EN ─────────────────────────────────────────────────────────

def _gen_docx_en(d: Dict, path: str, tmpdir: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    products_with_imgs = _decode_images(d["products"], tmpdir)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    if logo_path:
        try:
            for section in doc.sections:
                header = section.header
                hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = hp.add_run()
                run.add_picture(logo_path, height=Cm(1.2))
        except Exception:
            pass

    title = doc.add_heading("SALES CONTRACT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Contract No.: CTR-{d['date'].replace('-','')}-001")
    doc.add_paragraph(f"Date: {d['date']}")

    doc.add_heading("1. Parties", 1)
    doc.add_paragraph(f"Seller: {d['seller_name']}")
    doc.add_paragraph(f"Address: {d['seller_address']}")
    doc.add_paragraph(f"Contact: {d['seller_contact']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Buyer: {d['buyer_name']}")
    doc.add_paragraph(f"Address: {d['buyer_address']}")
    doc.add_paragraph(f"Contact: {d['buyer_contact']}")

    doc.add_heading("2. Products", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Product Name", "Product No./SKU", "Qty", "Unit Price (USD)", "Subtotal (USD)"]):
        hdr[i].text = h
    for p in d["products"]:
        row = table.add_row().cells
        for i, val in enumerate([p["name"], p["sku"], str(p["qty"]),
                                  _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])]):
            row[i].text = val
    ship_row = table.add_row().cells
    for i, val in enumerate(["Shipping", f"USD {_fmt_num(d['shipping_per'])} / unit", str(d["qty_total"]),
                              _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])]):
        ship_row[i].text = val
    total_row = table.add_row().cells
    total_row[3].text = "TOTAL"
    total_row[4].text = _fmt_num(d["grand_total"])

    doc.add_heading("3. Payment", 1)
    doc.add_paragraph("Buyer shall complete payment within 7 business days of contract signing. Accepted: USDT (Polygon) or bank wire.")

    doc.add_heading("4. Delivery", 1)
    doc.add_paragraph("Seller shall ship within 15 business days after receipt of full payment via DHL/FedEx.")

    doc.add_heading("5. Warranty", 1)
    doc.add_paragraph("12-month quality warranty from delivery date. Seller bears cost of defect remediation.")

    doc.add_heading("6. Liability", 1)
    doc.add_paragraph("Breaching party pays 10% of contract value as penalty plus actual damages.")

    doc.add_heading("7. Dispute Resolution", 1)
    doc.add_paragraph("Governed by the laws of Saint Lucia. Disputes submitted to agreed arbitration body.")

    doc.add_heading("8. Signatures", 1)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    sig_table.cell(0, 0).text = "Seller Signature / Stamp"
    sig_table.cell(0, 1).text = "Buyer Signature / Stamp"
    sig_table.cell(1, 0).text = "\n\n"
    sig_table.cell(1, 1).text = "\n\n"
    sig_table.cell(2, 0).text = f"Date: {d['date']}"
    sig_table.cell(2, 1).text = "Date: ___________"

    if d["needs_spec"]:
        doc.add_heading("Appendix: Product Specifications", 1)
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            doc.add_heading(label, 2)
            if spec_text:
                doc.add_paragraph(spec_text)
            for img_path in img_paths:
                try:
                    doc.add_picture(img_path, width=Cm(15))
                except Exception:
                    pass

    doc.save(path)


# ── PDF generator TW ──────────────────────────────────────────────────────────

def _gen_pdf_tw(d: Dict, path: str, tmpdir: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    _tw_font = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/PingFang.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJK_TW", font_path))
                _tw_font = "CJK_TW"
            except Exception:
                pass
            break

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=3.2*cm, bottomMargin=2*cm)
    title_style = ParagraphStyle("title_tw", fontName=_tw_font, fontSize=16,
                                 spaceAfter=6, alignment=1, leading=22)
    h2_style    = ParagraphStyle("h2_tw", fontName=_tw_font, fontSize=12,
                                 spaceAfter=4, spaceBefore=12, leading=18)
    h3_style    = ParagraphStyle("h3_tw", fontName=_tw_font, fontSize=11,
                                 spaceAfter=4, spaceBefore=8, leading=16)
    body_style  = ParagraphStyle("body_tw", fontName=_tw_font, fontSize=10,
                                 spaceAfter=4, leading=16)

    products_with_imgs = _decode_images(d["products"], tmpdir)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    header_cb = _make_header_cb(logo_path)

    story = []
    story.append(Paragraph("產品銷售合約", title_style))
    story.append(Paragraph(f"合約編號：CTR-{d['date'].replace('-','')}-001", body_style))
    story.append(Paragraph(f"簽訂日期：{d['date']}", body_style))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("一、甲乙雙方資訊", h2_style))
    story.append(Paragraph(f"<b>甲方（賣方）：</b>{d['seller_name']}", body_style))
    story.append(Paragraph(f"<b>地址：</b>{d['seller_address']}", body_style))
    story.append(Paragraph(f"<b>聯絡方式：</b>{d['seller_contact']}", body_style))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f"<b>乙方（買方）：</b>{d['buyer_name']}", body_style))
    story.append(Paragraph(f"<b>地址：</b>{d['buyer_address']}", body_style))
    story.append(Paragraph(f"<b>聯絡方式：</b>{d['buyer_contact']}", body_style))

    story.append(Paragraph("二、產品資訊", h2_style))
    tdata = [["產品名稱", "產品編號/SKU", "數量", "單價(USD)", "小計(USD)"]]
    for p in d["products"]:
        tdata.append([p["name"], p["sku"], str(p["qty"]),
                      _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])])
    tdata.append(["運費", f"每件 USD {_fmt_num(d['shipping_per'])}", str(d["qty_total"]),
                  _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])])
    tdata.append(["", "", "", "合計", _fmt_num(d["grand_total"])])

    t = Table(tdata, colWidths=[4.5*cm, 3*cm, 2*cm, 3*cm, 3*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,-1), _tw_font),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
        ("ALIGN",      (2,0), (-1,-1), "CENTER"),
        ("BACKGROUND", (0,-1), (-1,-1), colors.HexColor("#f0fdf4")),
        ("ROWBACKGROUNDS", (0,1), (-1,-2), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(t)

    story.append(Paragraph("三、付款方式", h2_style))
    story.append(Paragraph("乙方應於合約簽訂後 <b>7 個工作日內</b>完成付款，支援 USDT（Polygon 網路）或銀行電匯。", body_style))

    story.append(Paragraph("四、交貨條款", h2_style))
    story.append(Paragraph("甲方於收到全額貨款後 <b>15 個工作日內</b>安排出貨，運輸方式為國際快遞（DHL/FedEx）。", body_style))

    story.append(Paragraph("五、品質保證", h2_style))
    story.append(Paragraph("產品自交貨之日起享有 <b>12 個月</b>品質保固，因產品品質問題導致的損失由甲方承擔。", body_style))

    story.append(Paragraph("六、違約責任", h2_style))
    story.append(Paragraph("任何一方違約，應向守約方支付合約總金額 <b>10%</b> 的違約金，並賠償實際損失。", body_style))

    story.append(Paragraph("七、爭議解決", h2_style))
    story.append(Paragraph("本合約適用聖露西亞法律，爭議提交合約簽訂地仲裁委員會仲裁解決。", body_style))

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("八、簽署", h2_style))
    sig_data = [
        ["甲方（賣方）簽字/蓋章", "乙方（買方）簽字/蓋章"],
        ["\n\n\n", "\n\n\n"],
        [f"日期：{d['date']}", "日期：___________"],
    ]
    sig_t = Table(sig_data, colWidths=[8*cm, 8*cm])
    sig_t.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), _tw_font),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("BOX",      (0,0), (0,-1), 0.5, colors.HexColor("#334155")),
        ("BOX",      (1,0), (1,-1), 0.5, colors.HexColor("#334155")),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(sig_t)

    if d["needs_spec"]:
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph("附件：產品規格說明 / Product Specifications", h2_style))
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            story.append(Paragraph(label, h3_style))
            if spec_text:
                story.append(Paragraph(spec_text, body_style))
            for img_path in img_paths:
                try:
                    img = RLImage(img_path, width=15*cm)
                    img_w, img_h = img.imageWidth, img.imageHeight
                    if img_h > 0:
                        ratio = img_w / img_h
                        display_h = min(15*cm / ratio, 200)
                        img = RLImage(img_path, width=15*cm, height=display_h)
                    story.append(img)
                    story.append(Spacer(1, 0.3*cm))
                except Exception:
                    pass

    doc.build(story, onFirstPage=header_cb, onLaterPages=header_cb)


# ── Word generator TW ─────────────────────────────────────────────────────────

def _gen_docx_tw(d: Dict, path: str, tmpdir: str) -> None:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    products_with_imgs = _decode_images(d["products"], tmpdir)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    logo_path = _save_logo(d.get("logo_b64", ""), tmpdir)
    if logo_path:
        try:
            for section in doc.sections:
                header = section.header
                hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = hp.add_run()
                run.add_picture(logo_path, height=Cm(1.2))
        except Exception:
            pass

    title = doc.add_heading("產品銷售合約", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"合約編號：CTR-{d['date'].replace('-','')}-001")
    doc.add_paragraph(f"簽訂日期：{d['date']}")

    doc.add_heading("一、甲乙雙方資訊", 1)
    doc.add_paragraph(f"甲方（賣方）：{d['seller_name']}")
    doc.add_paragraph(f"地址：{d['seller_address']}")
    doc.add_paragraph(f"聯絡方式：{d['seller_contact']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"乙方（買方）：{d['buyer_name']}")
    doc.add_paragraph(f"地址：{d['buyer_address']}")
    doc.add_paragraph(f"聯絡方式：{d['buyer_contact']}")

    doc.add_heading("二、產品資訊", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["產品名稱", "產品編號/SKU", "數量", "單價(USD)", "小計(USD)"]):
        hdr[i].text = h
    for p in d["products"]:
        row = table.add_row().cells
        for i, val in enumerate([p["name"], p["sku"], str(p["qty"]),
                                  _fmt_num(p["unit_price"]), _fmt_num(p["subtotal"])]):
            row[i].text = val
    ship_row = table.add_row().cells
    for i, val in enumerate(["運費", f"每件 USD {_fmt_num(d['shipping_per'])}", str(d["qty_total"]),
                              _fmt_num(d["shipping_per"]), _fmt_num(d["shipping_total"])]):
        ship_row[i].text = val
    total_row = table.add_row().cells
    total_row[3].text = "合計"
    total_row[4].text = _fmt_num(d["grand_total"])

    doc.add_heading("三、付款方式", 1)
    doc.add_paragraph("乙方應於合約簽訂後 7 個工作日內完成付款，支援 USDT（Polygon 網路）或銀行電匯。")

    doc.add_heading("四、交貨條款", 1)
    doc.add_paragraph("甲方於收到全額貨款後 15 個工作日內安排出貨，運輸方式為國際快遞（DHL/FedEx）。")

    doc.add_heading("五、品質保證", 1)
    doc.add_paragraph("產品自交貨之日起享有 12 個月品質保固，因產品品質問題導致的損失由甲方承擔。")

    doc.add_heading("六、違約責任", 1)
    doc.add_paragraph("任何一方違約，應向守約方支付合約總金額 10% 的違約金，並賠償實際損失。")

    doc.add_heading("七、爭議解決", 1)
    doc.add_paragraph("本合約適用聖露西亞法律，爭議提交合約簽訂地仲裁委員會仲裁解決。")

    doc.add_heading("八、簽署", 1)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    sig_table.cell(0, 0).text = "甲方（賣方）簽字/蓋章"
    sig_table.cell(0, 1).text = "乙方（買方）簽字/蓋章"
    sig_table.cell(1, 0).text = "\n\n"
    sig_table.cell(1, 1).text = "\n\n"
    sig_table.cell(2, 0).text = f"日期：{d['date']}"
    sig_table.cell(2, 1).text = "日期：___________"

    if d["needs_spec"]:
        doc.add_heading("附件：產品規格說明 / Product Specifications", 1)
        for p in products_with_imgs:
            spec_text = p.get("spec_text", "")
            img_paths = p.get("_img_paths", [])
            if len(spec_text) <= 20 and not img_paths:
                continue
            label = f"[{p['sku']}] {p['name']}" if p.get("sku") else p["name"]
            doc.add_heading(label, 2)
            if spec_text:
                doc.add_paragraph(spec_text)
            for img_path in img_paths:
                try:
                    doc.add_picture(img_path, width=Cm(15))
                except Exception:
                    pass

    doc.save(path)


# ── public API ────────────────────────────────────────────────────────────────

def generate_contract(params: Dict) -> Dict[str, str]:
    """
    Generate contract files based on params.
    lang: "cn" | "tw" | "en"
    format: "pdf" | "docx" | "both"
    """
    d      = _build_data(params)
    lang   = params.get("lang", "cn")
    fmt    = params.get("format", "both")
    tmpdir = tempfile.mkdtemp(prefix="contract_")
    result: Dict[str, str] = {}

    want_pdf  = fmt in ("pdf", "both")
    want_docx = fmt in ("docx", "both")

    if lang == "cn":
        if want_pdf:
            p = os.path.join(tmpdir, "Contract_CN.pdf")
            _gen_pdf_cn(d, p, tmpdir)
            result["cn_pdf"] = p
        if want_docx:
            p = os.path.join(tmpdir, "Contract_CN.docx")
            _gen_docx_cn(d, p, tmpdir)
            result["cn_docx"] = p
    elif lang == "tw":
        if want_pdf:
            p = os.path.join(tmpdir, "Contract_TW.pdf")
            _gen_pdf_tw(d, p, tmpdir)
            result["tw_pdf"] = p
        if want_docx:
            p = os.path.join(tmpdir, "Contract_TW.docx")
            _gen_docx_tw(d, p, tmpdir)
            result["tw_docx"] = p
    elif lang == "en":
        if want_pdf:
            p = os.path.join(tmpdir, "Contract_EN.pdf")
            _gen_pdf_en(d, p, tmpdir)
            result["en_pdf"] = p
        if want_docx:
            p = os.path.join(tmpdir, "Contract_EN.docx")
            _gen_docx_en(d, p, tmpdir)
            result["en_docx"] = p

    return result
